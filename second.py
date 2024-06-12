#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed Jun 12 21:49:28 2024

@author: notsocertainwind
"""
import warnings
warnings.filterwarnings('ignore')

import numpy as np
import pandas as pd

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches

# Import the AreaConverter class from the modules.py file
from module import AreaConverter, AreaHelper

# Import specific methods
sqft_to_rapd = AreaConverter.sqft_to_rapd
sqm_to_rapd = AreaConverter.sqm_to_rapd
rapd_to_sqft = AreaConverter.rapd_to_sqft
rapd_to_sqm = AreaConverter.rapd_to_sqm
sqm_to_sqft = AreaConverter.sqm_to_sqft
sqft_to_sqm = AreaConverter.sqft_to_sqm


# ft2_to_m2 =AreaHelper.ft2_to_m2
# m2_to_ft2=AreaHelper.m2_to_ft2
# ft2_to_rapd=AreaHelper.ft2_to_rapd
# rapd_to_ft2=AreaHelper.rapd_to_ft2
# rapd_to_m2=AreaHelper.rapd_to_m2
# m2_to_rapd=AreaHelper.m2_to_rapd
# extract_rapd=AreaHelper.extract_rapd

# JSON-like data
data = {
    "name_client": ["Mr. Niraj Bhandari"],
    "address_client": ["Kanakai Ward No.04, Ilam"],
    "owner": ["Mr. Devi Prasad Bhandari"],
    "plot": ["111,222"],
    "area": ["224,553"],
    "address": ["kapan,dang"],
    "ctzno": ["641/3269"],
    "issuedate": ["2036/10/16 AD"],
    "address_owner": ["Kanakai Ward No.04, Ilam"],
    "relationship": ["Grandfather"],
    "final": ["4000000"]
}
# Create DataFrame
df_data = pd.DataFrame(data)



# Sample data for the table
data = {
    "name_client": ["Mr. Niraj Bhandari","Mr. Niraj Bhandari"],
    "address_client": ["Kanakai Ward No.04, Ilam","Kanakai Ward No.04, Ilam"],
    "owner": ["Mr. Devi Prasad Bhandari","Mr. Devi Prasad Bhandari"],
    "S.N.": [1, 2],
    "OWNERíS NAME": ["sameep", "sameep"],
    "LOCATION (ADDRESS OF PROPERTY)": ["kapan", "kapan"],
    "OWNERSHIP_TYPE": ["single", "single"],
    "PLOT NO.": [555, 666],
    "ctzno": ["641/3269","641/3269"],
    "issuedate": ["2036/10/16 AD","2036/10/16 AD"],
    "address_owner": ["Kanakai Ward No.04, Ilam","Kanakai Ward No.04, Ilam"],
    "relationship": ["Grandfather","Grandfather"],
    "AREA(m2)": [554.44, np.nan],
    "AREA(ft2)": [np.nan, 6032.15],  
    "AREA(rapd)": [np.nan, np.nan],
    'final':[8030730,8030730]
}



# Fill missing values with np.nan to ensure consistency in length
max_length = max(len(data["name_client"]), len(data["AREA(m2)"]), len(data["AREA(ft2)"]), len(data["AREA(rapd)"]))
for key in ["name_client", "AREA(m2)", "AREA(ft2)", "AREA(rapd)"]:
    if len(data[key]) < max_length:
        data[key] += [np.nan] * (max_length - len(data[key]))

# Convert to DataFrame
df = pd.DataFrame(data)
print(df)


# Convert to DataFrame
df = pd.DataFrame(data)


# ---------------------------------



# Define helper functions for each conversion
def ft2_to_m2(ft2):
    if pd.notnull(ft2):
        return sqft_to_sqm(ft2)
    else:
        return np.nan

def m2_to_ft2(m2):
    if pd.notnull(m2):
        return sqm_to_sqft(m2)
    else:
        return np.nan

def ft2_to_rapd(ft2):
    if pd.notnull(ft2):
        ropani, aana, paisa, daam = sqft_to_rapd(ft2)
        return f"{ropani}-{aana}-{paisa}-{daam}"
    else:
        return np.nan

def rapd_to_ft2(rapd):
    if pd.notnull(rapd):
        ropani, aana, paisa, daam = rapd
        return rapd_to_sqft(ropani, aana, paisa, daam)
    else:
        return np.nan

def rapd_to_m2(rapd):
    if pd.notnull(rapd):
        ropani, aana, paisa, daam = rapd
        return rapd_to_sqm(ropani, aana, paisa, daam)
    else:
        return np.nan

def m2_to_rapd(m2):
    if pd.notnull(m2):
        ft2 = sqm_to_sqft(m2)
        return ft2_to_rapd(ft2)
    else:
        return np.nan

# Extract rapd from the string representation for conversion
def extract_rapd(rapd_str):
    if pd.notnull(rapd_str):
        parts = rapd_str.split(", ")
        return tuple(float(part.split()[0]) for part in parts)
    else:
        return np.nan

# Fill missing values
for i, row in df.iterrows():
    if pd.isnull(row["AREA(m2)"]):
        if pd.notnull(row["AREA(ft2)"]):
            df.at[i, "AREA(m2)"] = ft2_to_m2(row["AREA(ft2)"])
        elif pd.notnull(row["AREA(rapd)"]):
            rapd = extract_rapd(row["AREA(rapd)"])
            df.at[i, "AREA(m2)"] = rapd_to_m2(rapd)
    
    if pd.isnull(row["AREA(ft2)"]):
        if pd.notnull(row["AREA(m2)"]):
            df.at[i, "AREA(ft2)"] = m2_to_ft2(row["AREA(m2)"])
        elif pd.notnull(row["AREA(rapd)"]):
            rapd = extract_rapd(row["AREA(rapd)"])
            df.at[i, "AREA(ft2)"] = rapd_to_ft2(rapd)
    
    if pd.isnull(row["AREA(rapd)"]):
        if pd.notnull(row["AREA(ft2)"]):
            df.at[i, "AREA(rapd)"] = ft2_to_rapd(row["AREA(ft2)"])
        elif pd.notnull(row["AREA(m2)"]):
            df.at[i, "AREA(rapd)"] = m2_to_rapd(row["AREA(m2)"])


#1 Ropani= 16 Aana 
# 1 Aana = 4 Paisa 
# 1 Paisa = 4 Daam 

def rapd_to_ropani(rapd):
    # Define conversion factors
    aana_per_ropani = 16
    paisa_per_aana = 4
    daam_per_paisa = 4

    # Split the rapd string into ropani, aana, paisa, and daam
    ropani, aana, paisa, daam = map(float, rapd.split("-"))

    # Convert paisa to aana
    total_aana = aana + (paisa / paisa_per_aana)

    # Convert aana to ropani
    total_ropani = ropani + (total_aana / aana_per_ropani)

    return total_ropani

print(df)
df['AREA(rapd)'].iloc[0]


final =df['final'][0]
list(df)

df['AREA(rapd)'].iloc[1]

df['Ropani']=df['AREA(rapd)'].apply(rapd_to_ropani)

list(df)
df['Ropani'].head()
total =df['Ropani'].sum()
print(total)

#Calculatiing rateAdopted,weightedAverage and marketRate
rateAdopted = round(final/total)
weightedAverage = round(rateAdopted +0.1*rateAdopted)
marketRate = round(weightedAverage + 0.1 *weightedAverage)

type(final)

#---------------------------
# Function to add a paragraph with bold data
def add_bold_paragraph(doc, text, bold_part):
    para = doc.add_paragraph()
    para.add_run(text).bold = False
    para.add_run(bold_part).bold = True

def set_cell_border(cell, **kwargs):
    """
    Set cell border for a table cell.
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edgeEl = OxmlElement(f'w:{edge}')
            edgeEl.set(qn('w:val'), kwargs.get(edge))
            edgeEl.set(qn('w:sz'), '4')  # 4 for 1/2 pt, 8 for 1 pt
            edgeEl.set(qn('w:space'), '0')
            edgeEl.set(qn('w:color'), '000000')  # black color
            tcPr.append(edgeEl)

# # Values to be inserted into the table
# market_rate = round(44.1892)
# weighted_average = round(40.172)
# rate_adopted = round(36.52)
# Creating the DataFrame
list(df)
df_table = df[['S.N.','OWNERíS NAME','LOCATION (ADDRESS OF PROPERTY)','AREA(rapd)','PLOT NO.']]
# df_table['Plot']=df_data['plot']
# Create a new Document
doc = Document()

# Extracting data
name_client = df.loc[0, "name_client"]
address_client = df.loc[0, "address_client"]
owner = df.loc[0, "owner"]
ctzno = df.loc[0, "ctzno"]
issuedate = df.loc[0, "issuedate"]
address_owner = df.loc[0, "address_owner"]
relationship = df.loc[0, "relationship"]
plot = df_data.loc[0, "plot"]

doc.add_heading('FIRST CARD', 0)
# Adding paragraphs with bold data
add_bold_paragraph(doc, "Name of Client/Applicant: ", name_client)
add_bold_paragraph(doc, "Address of Client: ", address_client)
add_bold_paragraph(doc, f"Owner of the Property (Plot no: {plot}): ", owner)
add_bold_paragraph(doc, "Citizenship No. & Issued Date: ", f"{ctzno}, {issuedate}")
add_bold_paragraph(doc, "Address of the Owner: ", address_owner)
add_bold_paragraph(doc, "Relationship with the Client: ", relationship)

doc.add_section()
doc.add_heading('DETAIL OF PROPERTY VALUED:', 1)




# Adding a table with the data
table = doc.add_table(rows=1, cols=len(df_table.columns))
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(df_table.columns):
    hdr_cells[i].text = column_name
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Adding data to the table
for index, row in df_table.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)
        for paragraph in row_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

# Function to set cell border
def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edgeEl = OxmlElement(f'w:{edge}')
            edgeEl.set(qn('w:val'), kwargs.get(edge))
            edgeEl.set(qn('w:sz'), '4')  # 4 for 1/2 pt, 8 for 1 pt
            edgeEl.set(qn('w:space'), '0')
            edgeEl.set(qn('w:color'), '000000')  # black color
            tcPr.append(edgeEl)

# Applying border to all cells
for row in table.rows:
    for cell in row.cells:
        set_cell_border(cell, top="single", left="single", bottom="single", right="single")

# Function to set column width
def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width

# # Setting different column widths
# column_widths = [Inches(1), Inches(4), Inches(8), Inches(4), Inches(3), Inches(1.5), Inches(2), Inches(2.5)]

# for i, column in enumerate(table.columns):
#     set_column_width(column, column_widths[i])
doc.add_section()
doc.add_heading('VALUE OF LAND', 0)

# Add a table with 2 columns and 3 rows
table = doc.add_table(rows=3, cols=2)
# Set table style
table.style = 'Table Grid'

# Set bold font style
bold_style = doc.styles['Normal']
bold_style.font.bold = True

# Nepali numbers  format as "44,37,070"
formatted_marketRate = "{:,.0f}".format(marketRate)
formatted_weightedAverage = "{:,.0f}".format(weightedAverage)
formatted_rateAdopted = "{:,.0f}".format(rateAdopted)

type(formatted_rateAdopted)
# Populate the table
# Populate the table
table.cell(0, 0).text = "Market rate *(Current Market Rate inquired from Local Resident (RS.)/Ropani )"
table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
table.cell(0, 1).text = formatted_marketRate
table.cell(1, 0).text = "Weighted Average Calculation:"
table.cell(1, 0).paragraphs[0].runs[0].font.bold = True
table.cell(1, 1).text = formatted_weightedAverage
table.cell(2, 0).text = "Land Rate Adopted in this Report (Rs.)/Ropani:"
table.cell(2, 0).paragraphs[0].runs[0].font.bold = True
table.cell(2, 1).text = formatted_rateAdopted

# Applying border to all cells
for row in table.rows:
    for cell in row.cells:
        set_cell_border(cell, top="single", left="single", bottom="single", right="single")



doc.add_section()

#-------------------------



# Sample data
data_three = {
    'Description': ['Regular', 'Regular', 'Total']
}
list(df)
df.shape

df_three = pd.DataFrame(data_three)

df_three['Plot No.']=df['PLOT NO.']
df_three['Plot Area (sqm.)']=round(df['AREA(m2)'],2)
df_three['Area in Sq.ft.']=round(df['AREA(ft2)'],2)
df_three['Area in Ropani.']=round(df['Ropani'],3)
df_three['Area in (R-A-P-D)']=df['AREA(rapd)']

df_three
list(df_three)
print(df_three.iloc[0])
print(df_three.iloc[1])
print(df_three.iloc[2])

# Summing the numeric columns for the third row

df_three.at[2, 'Plot Area (sqm.)'] = round(df_three['Plot Area (sqm.)'][:-1].sum(),2)
df_three.at[2, 'Area in Sq.ft.'] = round(df_three['Area in Sq.ft.'][:-1].sum(),2)
df_three.at[2, 'Area in Ropani.'] = round(df_three['Area in Ropani.'][:-1].sum(),2)
# Function to sum the parts of 'Area in (R-A-P-D)'
def sum_r_a_p_d(values):
    r, a, p, d = 0, 0, 0, 0.0
    for value in values:
        parts = value.split('-')
        r += int(parts[0])
        a += int(parts[1])
        p += int(parts[2])
        d += round(float(parts[3]))
    return f'{r}-{a}-{p}-{d:.3f}'

df_three.at[2, 'Area in (R-A-P-D)'] = sum_r_a_p_d(df_three['Area in (R-A-P-D)'][:-1])


print(df_three.iloc[2])
def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "FF0000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "00FF00", "space": "0"},
        start={"sz": 24, "val": "dashed", "color": "0000FF", "space": "0"},
        end={"sz": 12, "val": None, "color": "000000", "space": "0"},
    )
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edgeEl = OxmlElement(f'w:{edge}')
            edgeEl.set(qn('w:val'), kwargs.get(edge))
            edgeEl.set(qn('w:sz'), '4')  # 4 for 1/2 pt, 8 for 1 pt
            edgeEl.set(qn('w:space'), '0')
            edgeEl.set(qn('w:color'), '000000')  # black color
            tcPr.append(edgeEl)

def create_table(doc, title, df):
    doc.add_paragraph(title, style='Heading 1')
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        set_cell_border(hdr_cells[i], top="single", left="single", bottom="single", right="single")

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value) if not pd.isna(value) else ''
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_border(row_cells[i], top="single", left="single", bottom="single", right="single")

    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = 1  # center align


doc.add_heading('Plot Area Summary', 0)


# Create the first table "As per Lalpurja"
create_table(doc, "As per Lalpurja", df_three)

# Create the second table "As per Measurement"
create_table(doc, "As per Measurement", df_three)

# Create the third table "The minimum area is taken for valuation"
create_table(doc, "The minimum area is taken for valuation", df_three)



#----------

# Save the document
file_path = "Land_Rate_Report_with_Borders.docx"
doc.save(file_path)

print(f"Document saved at: {file_path}")


plot
print(list(df))
print(final)
print(rateAdopted)
print(total)
list(df_three)
ropani_total =df_three['Area in (R-A-P-D)'].iloc[-1]
ropani_total

# # weightedAverage  marketRate


# # Function to sum the parts of 'Area in (R-A-P-D)'
# def sum_r_a_p_d(values):
#     r, a, p, d = 0, 0, 0, 0.0
#     for value in values:
#         parts = value.split('-')
#         r += int(parts[0])
#         a += int(parts[1])
#         p += int(parts[2])
#         d += float(parts[3])
#     return f'{r}-{a}-{p}-{d:.3f}'

# df.at[2, 'Area in (R-A-P-D)'] = sum_r_a_p_d(df['Area in (R-A-P-D)'][:-1])

